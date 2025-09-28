from __future__ import annotations

from pathlib import Path
from typing import List
import docx  # python-docx

from .models import ParseOutput, ParsedTable


def parse_docx(path: Path) -> ParseOutput:
    document = docx.Document(str(path))

    tables: List[ParsedTable] = []
    for t in document.tables:
        rows: List[List[str]] = []
        headers: List[str] = []
        if t.rows:
            # header as first row if non-empty
            raw_header = [cell.text.strip() for cell in t.rows[0].cells]
            if any(h for h in raw_header):
                headers = raw_header
                data_rows = t.rows[1:]
            else:
                data_rows = t.rows
            for r in data_rows:
                rows.append([cell.text.strip() for cell in r.cells])
        tables.append(ParsedTable(headers=headers, rows=rows, meta={}))

    # paragraph text (single page output)
    para_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text and p.text.strip())
    pages_text = [para_text] if para_text else []

    return ParseOutput(source_path=path, pages_text=pages_text, tables=tables)


